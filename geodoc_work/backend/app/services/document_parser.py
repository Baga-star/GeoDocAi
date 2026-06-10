"""
Hybrid structure-preserving parser for geological documents.

The parser returns typed artifacts instead of flat text chunks:
text, table, figure, map. Tables are preserved with columns/rows; captions for
figures/maps are indexed separately. OCR is used only when the page/file has no
usable text layer or when an image is uploaded.
"""
import asyncio
import csv
import io
import json
import logging
import re
from pathlib import Path
from typing import Any

from fastapi import UploadFile

from app.models import DocumentArtifact
from app.services.local_index import chunk_text, table_to_markdown
from app.services.text_cleanup import clean_cell_text, looks_like_intrusive_noise, normalize_pdf_prose
from app.services.vision_agent import (
    analyze_geology_image,
    image_data_url,
    vision_is_configured,
    visual_payload_to_text,
    visual_payload_title,
    visual_payload_type,
)


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls",
    ".csv", ".txt", ".md",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp",
}

CAPTION_RE = re.compile(
    r"(?im)^\s*((?:Таблица|Рисунок|Рис\.|Карта|Схема|Figure|Fig\.|Map)\s*\d+(?:[\.\-]\d+)?[^\n]*)"
)
TABLE_TITLE_RE = re.compile(r"(?im)^\s*(Таблица\s*\d+(?:[\.\-]\d+)?[^\n]*)")
FIGURE_TITLE_RE = re.compile(r"(?im)^\s*((?:Рисунок|Рис\.|Схема|Figure|Fig\.)\s*\d+(?:[\.\-]\d+)?[^\n]*)")
MAP_TITLE_RE = re.compile(r"(?im)^\s*((?:Карта|Map)\s*\d+(?:[\.\-]\d+)?[^\n]*)")
UNIT_RE = re.compile(r"\b(?:т/сут|м3/сут|м³/сут|МПа|кг/см²|кг/см2|г/см³|г/см3|тыс\.\s*т|м)\b", re.I)


async def parse_upload(file: UploadFile) -> list[DocumentArtifact]:
    filename = file.filename or "document"
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Формат «{suffix or 'без расширения'}» не поддерживается. "
            "Поддерживаются: PDF, DOCX, Excel, CSV, TXT/MD и изображения."
        )

    content = await file.read()
    if not content:
        raise ValueError("Файл пустой или не удалось прочитать содержимое.")

    if suffix == ".pdf":
        artifacts = await _parse_pdf(filename, content)
    elif suffix in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".webp"}:
        artifacts = await _parse_image(filename, content, suffix)
    elif suffix == ".docx":
        artifacts = _parse_docx(filename, content)
    elif suffix in {".xlsx", ".xls"}:
        artifacts = _parse_excel(filename, content)
    elif suffix == ".csv":
        artifacts = _parse_csv(filename, content)
    else:
        text = content.decode("utf-8", errors="ignore")
        artifacts = _text_to_artifacts(filename, text)

    if not artifacts:
        raise ValueError(
            f"Не удалось извлечь содержимое из «{filename}». "
            "Если это скан или изображение, укажите DEEPSEEK_OCR_API_KEY в .env."
        )

    logger.info("Файл '%s': извлечено %d artifacts", filename, len(artifacts))
    return artifacts


# ── PDF ───────────────────────────────────────────────────────────────────
async def _parse_pdf(filename: str, content: bytes) -> list[DocumentArtifact]:
    # Run CPU-bound pdfplumber extraction in a thread pool so it doesn't block the event loop
    artifacts = await asyncio.to_thread(_parse_pdf_with_pdfplumber, filename, content)

    # Visual pages are first-class evidence. Scanned maps often have no text
    # layer, so they must be indexed even before OCR succeeds.
    from app.config import get_settings
    settings = get_settings()
    visual_artifacts = await _parse_pdf_visual_pages(filename, content, artifacts)
    artifacts = _merge_visual_artifacts(artifacts, visual_artifacts)

    usable = _has_usable_artifacts(artifacts)
    if usable:
        return artifacts

    # Scanned PDF or broken text layer: use OCR when configured.
    if not settings.deepseek_ocr_api_key or settings.deepseek_ocr_api_key == "replace-me":
        return artifacts

    try:
        from app.services.ocr_client import ocr_pdf_bytes
        raw = await ocr_pdf_bytes(
            content,
            api_key=settings.deepseek_ocr_api_key,
            api_url=settings.deepseek_ocr_url,
            model=settings.deepseek_ocr_model,
            max_pages=settings.ocr_max_pages,
            structured=True,
        )
        ocr_artifacts = _ocr_output_to_artifacts(filename, raw)
        return _merge_visual_artifacts(ocr_artifacts, visual_artifacts) or artifacts
    except Exception as exc:
        logger.warning("DeepSeek OCR PDF error: %s", exc)
        return artifacts



def _extract_pdf_page_texts_with_fitz(content: bytes) -> dict[int, str]:
    """PyMuPDF often preserves prose order better than pdfplumber on stamped PDFs."""
    try:
        import fitz
    except Exception:
        return {}
    texts: dict[int, str] = {}
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        try:
            for page_num in range(1, len(doc) + 1):
                raw = doc[page_num - 1].get_text("text") or ""
                clean = normalize_pdf_prose(raw)
                if clean:
                    texts[page_num] = clean
        finally:
            doc.close()
    except Exception as exc:
        logger.debug("fitz text extraction failed: %s", exc)
    return texts

def _best_page_text(pdfplumber_text: str, fitz_text: str | None) -> str:
    plumber = normalize_pdf_prose(pdfplumber_text or "")
    fitz_clean = normalize_pdf_prose(fitz_text or "")
    if not fitz_clean:
        return plumber
    if not plumber:
        return fitz_clean
    # Prefer PyMuPDF when pdfplumber contains watermark-letter contamination.
    if looks_like_intrusive_noise(plumber, threshold=0.12) and not looks_like_intrusive_noise(fitz_clean, threshold=0.18):
        return fitz_clean
    # If both are clean, use the longer one, but avoid a huge table-only blob.
    if len(fitz_clean) > len(plumber) * 1.15:
        return fitz_clean
    return plumber

def _parse_pdf_with_pdfplumber(filename: str, content: bytes) -> list[DocumentArtifact]:
    try:
        import pdfplumber
    except ImportError:
        logger.warning("pdfplumber не установлен")
        return _parse_pdf_with_pypdf(filename, content)

    artifacts: list[DocumentArtifact] = []
    fitz_texts = _extract_pdf_page_texts_with_fitz(content)
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                plumber_text_raw = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
                page_text = _best_page_text(plumber_text_raw, fitz_texts.get(page_num))
                table_objects = []
                try:
                    table_objects = page.find_tables()
                except Exception as exc:
                    logger.debug("find_tables failed on page %s: %s", page_num, exc)

                table_bboxes: list[tuple[float, float, float, float]] = []
                used_titles: set[str] = set()
                for table_idx, table_obj in enumerate(table_objects, start=1):
                    try:
                        raw_table = table_obj.extract()
                    except Exception:
                        raw_table = None
                    if not raw_table:
                        continue
                    table_bboxes.append(tuple(float(v) for v in table_obj.bbox))
                    columns, rows = _clean_table(raw_table)
                    if not columns or not rows:
                        continue
                    title = _nearest_table_title(page_text, used_titles) or f"Таблица на стр. {page_num}"
                    used_titles.add(title)
                    markdown = table_to_markdown(columns, rows)
                    text = f"{title}\n{markdown}"
                    artifacts.append(
                        DocumentArtifact(
                            document_name=filename,
                            page=page_num,
                            artifact_type="table",
                            title=title,
                            text=text,
                            columns=columns,
                            rows=rows,
                            units=_extract_units(text),
                            bbox=[float(v) for v in table_obj.bbox],
                            metadata={"table_index": table_idx, "parser": "pdfplumber"},
                        )
                    )

                # Visual captions must become separate searchable artifacts.
                artifacts.extend(_caption_artifacts(filename, page_num, page_text))

                # Extract text outside tables. This avoids duplicating table gibberish.
                text_outside_tables = _extract_text_outside_bboxes(page, table_bboxes) or page_text
                text_outside_tables = _best_page_text(text_outside_tables, fitz_texts.get(page_num) or page_text)
                text_outside_tables = normalize_pdf_prose(_remove_known_captions(text_outside_tables))
                for chunk in chunk_text(text_outside_tables):
                    if _chunk_is_mostly_table_noise(chunk):
                        continue
                    artifacts.append(
                        DocumentArtifact(
                            document_name=filename,
                            page=page_num,
                            artifact_type="text",
                            text=chunk,
                            units=_extract_units(chunk),
                            metadata={"parser": "pdfplumber"},
                        )
                    )
        return artifacts
    except Exception as exc:
        logger.warning("pdfplumber parse error: %s", exc)
        return _parse_pdf_with_pypdf(filename, content)


def _parse_pdf_with_pypdf(filename: str, content: bytes) -> list[DocumentArtifact]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return []
    artifacts: list[DocumentArtifact] = []
    try:
        reader = PdfReader(io.BytesIO(content))
        for page_num, page in enumerate(reader.pages, start=1):
            text = normalize_pdf_prose(page.extract_text() or "")
            artifacts.extend(_caption_artifacts(filename, page_num, text))
            for chunk in chunk_text(_remove_known_captions(text)):
                artifacts.append(
                    DocumentArtifact(
                        document_name=filename,
                        page=page_num,
                        artifact_type="text",
                        text=chunk,
                        units=_extract_units(chunk),
                        metadata={"parser": "pypdf"},
                    )
                )
        return artifacts
    except Exception as exc:
        logger.warning("pypdf parse error: %s", exc)
        return []


# ── Images / OCR ───────────────────────────────────────────────────────────
async def _parse_image(filename: str, content: bytes, suffix: str) -> list[DocumentArtifact]:
    from app.config import get_settings
    settings = get_settings()
    mime_type = _mime_from_suffix(suffix)

    # Always index images as visual artifacts so upload does not fail just
    # because the vision/OCR key has not been configured yet. When a vision
    # model is available, this artifact contains the actual visual reading.
    visual = await _visual_artifact_from_image_bytes(
        filename,
        content,
        page=1,
        mime_type=mime_type,
        fallback_type=_guess_visual_type(content),
        source="image-upload",
    )
    artifacts: list[DocumentArtifact] = [visual]

    # DeepSeek OCR can additionally recover scanned tables/text. Keep it as an
    # optional enrichment layer, not as a blocker for image uploads.
    if settings.deepseek_ocr_api_key and settings.deepseek_ocr_api_key != "replace-me":
        try:
            from app.services.ocr_client import ocr_image_bytes
            raw = await ocr_image_bytes(
                content,
                api_key=settings.deepseek_ocr_api_key,
                api_url=settings.deepseek_ocr_url,
                model=settings.deepseek_ocr_model,
                mime_type=mime_type,
                structured=True,
            )
            artifacts = _merge_visual_artifacts(_ocr_output_to_artifacts(filename, raw), [visual]) or artifacts
        except Exception as exc:
            logger.warning("Image OCR enrichment failed for %s: %s", filename, exc)
    return artifacts


# ── Office formats ────────────────────────────────────────────────────────
def _parse_docx(filename: str, content: bytes) -> list[DocumentArtifact]:
    try:
        from docx import Document
    except ImportError:
        return []
    artifacts: list[DocumentArtifact] = []
    try:
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for chunk in chunk_text("\n\n".join(paragraphs)):
            artifacts.append(DocumentArtifact(document_name=filename, artifact_type="text", text=chunk, units=_extract_units(chunk)))
        for table_idx, table in enumerate(doc.tables, start=1):
            raw = [[cell.text for cell in row.cells] for row in table.rows]
            columns, rows = _clean_table(raw)
            if columns and rows:
                title = _nearest_table_title("\n".join(paragraphs), set()) or f"Таблица DOCX {table_idx}"
                md = table_to_markdown(columns, rows)
                artifacts.append(
                    DocumentArtifact(
                        document_name=filename,
                        artifact_type="table",
                        title=title,
                        text=f"{title}\n{md}",
                        columns=columns,
                        rows=rows,
                        units=_extract_units(md),
                        metadata={"table_index": table_idx, "parser": "python-docx"},
                    )
                )
        return artifacts
    except Exception as exc:
        logger.warning("docx parse error: %s", exc)
        return []


def _parse_excel(filename: str, content: bytes) -> list[DocumentArtifact]:
    try:
        import pandas as pd
    except ImportError:
        return []
    artifacts: list[DocumentArtifact] = []
    try:
        workbook = pd.read_excel(io.BytesIO(content), sheet_name=None)
        for sheet_name, frame in workbook.items():
            frame = frame.fillna("")
            columns = [str(c).strip() for c in frame.columns]
            rows = [[str(v).strip() for v in row] for row in frame.values.tolist()]
            title = f"Лист Excel: {sheet_name}"
            md = table_to_markdown(columns, rows)
            artifacts.append(
                DocumentArtifact(
                    document_name=filename,
                    artifact_type="table",
                    title=title,
                    text=f"{title}\n{md}",
                    columns=columns,
                    rows=rows,
                    units=_extract_units(md),
                    metadata={"sheet": sheet_name, "parser": "pandas"},
                )
            )
        return artifacts
    except Exception as exc:
        logger.warning("excel parse error: %s", exc)
        return []


def _parse_csv(filename: str, content: bytes) -> list[DocumentArtifact]:
    decoded = content.decode("utf-8-sig", errors="ignore")
    rows = list(csv.reader(io.StringIO(decoded)))
    columns, body = _clean_table(rows)
    if not columns:
        return _text_to_artifacts(filename, decoded)
    title = f"CSV таблица: {filename}"
    md = table_to_markdown(columns, body)
    return [
        DocumentArtifact(
            document_name=filename,
            artifact_type="table",
            title=title,
            text=f"{title}\n{md}",
            columns=columns,
            rows=body,
            units=_extract_units(md),
            metadata={"parser": "csv"},
        )
    ]


# ── Helpers ───────────────────────────────────────────────────────────────
def _text_to_artifacts(filename: str, text: str, page: int | None = None) -> list[DocumentArtifact]:
    artifacts = _caption_artifacts(filename, page, text)
    for chunk in chunk_text(_remove_known_captions(text)):
        artifacts.append(DocumentArtifact(document_name=filename, page=page, artifact_type="text", text=chunk, units=_extract_units(chunk)))
    return artifacts


def _has_usable_artifacts(artifacts: list[DocumentArtifact]) -> bool:
    if any(a.artifact_type == "table" and a.rows for a in artifacts):
        return True
    if any(a.artifact_type in {"figure", "map"} for a in artifacts):
        return True
    text_len = sum(len(a.text or "") for a in artifacts if a.artifact_type == "text")
    return text_len > 200


def _clean_cell(value: Any) -> str:
    return clean_cell_text(value)


def _clean_table(raw_table: list[list[Any]]) -> tuple[list[str], list[list[str]]]:
    cleaned = [[_clean_cell(cell) for cell in row] for row in raw_table if row]
    cleaned = [row for row in cleaned if any(cell for cell in row)]
    if not cleaned:
        return [], []
    width = max(len(row) for row in cleaned)
    cleaned = [(row + [""] * width)[:width] for row in cleaned]
    header_idx = 0
    # Skip empty/numeric row indexes if needed.
    for idx, row in enumerate(cleaned[:3]):
        non_empty = [c for c in row if c]
        if len(non_empty) >= 2:
            header_idx = idx
            break
    columns = cleaned[header_idx]
    if len(set(c for c in columns if c)) < len([c for c in columns if c]):
        columns = [f"{c or 'Колонка'} {i+1}".strip() for i, c in enumerate(columns)]
    columns = [c or f"Колонка {i+1}" for i, c in enumerate(columns)]
    rows = cleaned[header_idx + 1 :]
    values = [*columns, *[cell for row in rows[:10] for cell in row]]
    noisy = sum(1 for value in values if value and looks_like_intrusive_noise(value, threshold=0.16))
    if values and noisy / max(len(values), 1) > 0.10:
        # Do not persist fake table cells made from diagonal watermark letters.
        return [], []
    return columns, rows


def _nearest_table_title(page_text: str, used_titles: set[str]) -> str | None:
    for match in TABLE_TITLE_RE.findall(page_text or ""):
        title = re.sub(r"\s+", " ", match).strip()
        if title not in used_titles:
            return title
    return None


def _caption_artifacts(filename: str, page: int | None, text: str) -> list[DocumentArtifact]:
    artifacts: list[DocumentArtifact] = []
    seen: set[str] = set()
    for regex, artifact_type in ((MAP_TITLE_RE, "map"), (FIGURE_TITLE_RE, "figure")):
        for match in regex.findall(text or ""):
            title = re.sub(r"\s+", " ", match).strip()
            if not title or title in seen:
                continue
            seen.add(title)
            artifacts.append(
                DocumentArtifact(
                    document_name=filename,
                    page=page,
                    artifact_type=artifact_type,  # type: ignore[arg-type]
                    title=title,
                    caption=title,
                    text=title,
                    metadata={"parser": "caption-detector"},
                )
            )
    return artifacts


def _remove_known_captions(text: str) -> str:
    return CAPTION_RE.sub("", text or "").strip()


def _extract_units(text: str) -> list[str]:
    units: list[str] = []
    for unit in UNIT_RE.findall(text or ""):
        if unit not in units:
            units.append(unit)
    return units


def _extract_text_outside_bboxes(page: Any, bboxes: list[tuple[float, float, float, float]]) -> str:
    if not bboxes:
        return page.extract_text(x_tolerance=1, y_tolerance=3) or ""

    def outside_tables(obj: dict[str, Any]) -> bool:
        x0 = obj.get("x0")
        x1 = obj.get("x1")
        top = obj.get("top")
        bottom = obj.get("bottom")
        if x0 is None or x1 is None or top is None or bottom is None:
            return True
        cx = (float(x0) + float(x1)) / 2
        cy = (float(top) + float(bottom)) / 2
        for bx0, btop, bx1, bbottom in bboxes:
            if bx0 <= cx <= bx1 and btop <= cy <= bbottom:
                return False
        return True

    try:
        return page.filter(outside_tables).extract_text(x_tolerance=1, y_tolerance=3) or ""
    except Exception:
        return page.extract_text(x_tolerance=1, y_tolerance=3) or ""


def _chunk_is_mostly_table_noise(chunk: str) -> bool:
    lines = [l for l in chunk.splitlines() if l.strip()]
    if len(lines) < 2:
        return False
    short_numeric = sum(1 for l in lines if len(l) < 80 and len(re.findall(r"\d", l)) >= 3)
    return short_numeric / max(len(lines), 1) > 0.75


def _mime_from_suffix(suffix: str) -> str:
    return {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }.get(suffix.lower(), "image/png")


def _ocr_output_to_artifacts(filename: str, raw: str) -> list[DocumentArtifact]:
    raw = raw.strip()
    if not raw:
        return []
    try:
        data = json.loads(_strip_code_fence(raw))
    except Exception:
        return _text_to_artifacts(filename, raw)

    pages = data if isinstance(data, list) else [data]
    artifacts: list[DocumentArtifact] = []
    for page_obj in pages:
        if not isinstance(page_obj, dict):
            continue
        page = _safe_int(page_obj.get("page"))
        for block in page_obj.get("texts", []) or []:
            if isinstance(block, dict) and block.get("text"):
                artifacts.extend(_text_to_artifacts(filename, str(block.get("text")), page=page))
        for table in page_obj.get("tables", []) or []:
            if not isinstance(table, dict):
                continue
            columns = [str(c).strip() for c in (table.get("columns") or [])]
            rows = [[str(c).strip() for c in row] for row in (table.get("rows") or []) if isinstance(row, list)]
            title = table.get("title") or f"OCR таблица на стр. {page or '?'}"
            md = table_to_markdown(columns, rows)
            artifacts.append(
                DocumentArtifact(
                    document_name=filename,
                    page=page,
                    artifact_type="table",
                    title=str(title),
                    text=f"{title}\n{md}" if md else str(title),
                    columns=columns,
                    rows=rows,
                    units=[str(u) for u in (table.get("units") or [])] or _extract_units(md),
                    metadata={"parser": "deepseek-ocr", "notes": table.get("notes")},
                )
            )
        for key, typ in (("figures", "figure"), ("maps", "map")):
            for item in page_obj.get(key, []) or []:
                if not isinstance(item, dict):
                    continue
                title = item.get("title") or item.get("caption") or f"OCR {typ} на стр. {page or '?'}"
                caption = item.get("caption") or item.get("description") or title
                description = item.get("description") or item.get("summary") or ""
                text = "\n".join(str(v) for v in (title, caption, description) if v)
                artifacts.append(
                    DocumentArtifact(
                        document_name=filename,
                        page=page,
                        artifact_type=typ,  # type: ignore[arg-type]
                        title=str(title),
                        caption=str(caption),
                        text=text,
                        metadata={"parser": "deepseek-ocr", "visual_analysis": item},
                    )
                )
    return artifacts



# ── Visual map / figure reading ────────────────────────────────────────────
async def _parse_pdf_visual_pages(
    filename: str,
    content: bytes,
    existing_artifacts: list[DocumentArtifact],
) -> list[DocumentArtifact]:
    """Render selected PDF pages and run Visual Geology Reading Agent."""
    from app.config import get_settings
    settings = get_settings()
    candidates = _candidate_visual_pages(existing_artifacts, settings.vision_max_pages)
    if not candidates and not existing_artifacts:
        # Scanned/image-only PDF: every page can be a visual artifact.
        candidates = None
    elif not candidates:
        return []

    try:
        import fitz  # pymupdf
    except ImportError:
        logger.warning("pymupdf is not installed; visual PDF analysis is disabled")
        return []

    doc = fitz.open(stream=content, filetype="pdf")
    try:
        max_visual_pages = max(1, int(settings.vision_max_pages or 1))
        if candidates is None:
            page_numbers = list(range(1, min(len(doc), max_visual_pages) + 1))
        else:
            # `vision_max_pages` is a count limit, not a page-number ceiling.
            # Earlier builds skipped figures on pages like 44/70 because they were
            # greater than the max count. Keep selected pages anywhere in the PDF.
            page_numbers = [p for p in sorted(candidates) if 1 <= p <= len(doc)][:max_visual_pages]

        # Pre-render all pages to PNG in the thread pool (CPU-bound).
        def _render_pages() -> list[tuple[int, bytes]]:
            rendered = []
            for pn in page_numbers:
                pg = doc[pn - 1]
                matrix = fitz.Matrix(settings.vision_render_dpi / 72, settings.vision_render_dpi / 72)
                pix = pg.get_pixmap(matrix=matrix, alpha=False)
                rendered.append((pn, pix.tobytes("png")))
            return rendered

        rendered_pages = await asyncio.to_thread(_render_pages)
    finally:
        doc.close()

    # WHY: previously awaited each page vision-call one by one.
    # With 8 pages x ~30s each = 240s. Parallel gather cuts this to ~30s total.
    async def _analyze_one(page_num: int, png_bytes: bytes) -> DocumentArtifact:
        fallback_type = _fallback_type_for_page(existing_artifacts, page_num) or _guess_visual_type(png_bytes)
        return await _visual_artifact_from_image_bytes(
            filename,
            png_bytes,
            page=page_num,
            mime_type="image/png",
            fallback_type=fallback_type,
            source="pdf-rendered-page",
        )

    artifacts: list[DocumentArtifact] = list(
        await asyncio.gather(*[_analyze_one(pn, pb) for pn, pb in rendered_pages])
    )
    return artifacts


def _candidate_visual_pages(artifacts: list[DocumentArtifact], max_pages: int) -> set[int] | None:
    if not artifacts:
        return None
    pages: set[int] = set()
    by_page_text_len: dict[int, int] = {}
    for artifact in artifacts:
        if artifact.page is None:
            continue
        by_page_text_len[artifact.page] = by_page_text_len.get(artifact.page, 0) + len(artifact.text or "")
        hay = " ".join([artifact.artifact_type, artifact.title or "", artifact.caption or "", artifact.text or ""]).lower()
        if artifact.artifact_type in {"map", "figure"} or any(token in hay for token in ("карта", "рисунок", "рис.", "схема", "разрез", "map", "figure", "fig.")):
            pages.add(artifact.page)
    # Pages with almost no text are often scanned visual pages.
    for page, length in by_page_text_len.items():
        if length < 80:
            pages.add(page)
    return set(sorted(pages)[:max_pages])


def _fallback_type_for_page(artifacts: list[DocumentArtifact], page: int) -> str | None:
    for artifact in artifacts:
        if artifact.page != page:
            continue
        hay = " ".join([artifact.artifact_type, artifact.title or "", artifact.caption or "", artifact.text or ""]).lower()
        if artifact.artifact_type == "map" or "карт" in hay or "map" in hay:
            return "map"
        if artifact.artifact_type == "figure" or any(token in hay for token in ("рис", "схем", "разрез", "figure", "fig")):
            return "figure"
    return None


async def _visual_artifact_from_image_bytes(
    filename: str,
    image_bytes: bytes,
    *,
    page: int | None,
    mime_type: str,
    fallback_type: str,
    source: str,
) -> DocumentArtifact:
    from app.config import get_settings
    settings = get_settings()
    preview_url = _make_preview_data_url(image_bytes, mime_type)
    # Local CV heuristic is intentionally evaluated for every visual artifact.
    # Vision/OCR models sometimes describe a scanned structural map as a
    # generic "figure" or "scheme". GeoDoc must still route such pages to
    # Map Workspace when the image itself has map-like visual grammar
    # (colored reserve zones, contours/isolines, well points, legend blocks).
    local_visual_type = _guess_visual_type(image_bytes)
    local_map_reading = _local_map_reading(image_bytes) if local_visual_type == "map" or fallback_type == "map" else {}
    payload: dict[str, Any] | None = None
    parser = "visual-placeholder"

    if vision_is_configured(settings):
        try:
            payload = await analyze_geology_image(
                image_bytes,
                settings=settings,
                filename=filename,
                page=page,
                mime_type=mime_type,
            )
            parser = "visual-geology-agent"
        except Exception as exc:
            logger.warning("Visual analysis failed for %s page %s: %s", filename, page, exc)

    if payload:
        artifact_type = visual_payload_type(payload, fallback=fallback_type)
        if local_visual_type == "map" and artifact_type != "map":
            artifact_type = "map"
            payload = dict(payload)
            payload["visual_type"] = "map"
            observations = payload.get("observations")
            if not isinstance(observations, list):
                observations = [str(observations)] if observations else []
            observations.append("MapReader определил страницу как карту по геометрии: цветовые зоны, контуры/изолинии, точки и легенда.")
            payload["observations"] = observations
        if artifact_type == "map" and local_map_reading:
            payload = _merge_local_map_reading(payload, local_map_reading)
        text = visual_payload_to_text(payload)
        title = visual_payload_title(payload, artifact_type, filename, page)
        caption = payload.get("caption") or payload.get("title") or title
        metadata = {
            "parser": parser,
            "source": source,
            "visual_analysis": payload,
            "visual_status": "analyzed",
            "visual_model": settings.vision_model,
            "local_visual_type": local_visual_type,
            "local_map_reading": local_map_reading,
            "preview_data_url": preview_url,
        }
        if payload.get("confidence"):
            metadata["visual_confidence"] = payload.get("confidence")
    else:
        artifact_type = local_visual_type if local_visual_type in {"map", "figure"} else (fallback_type if fallback_type in {"map", "figure"} else "figure")
        title = _fallback_visual_title(artifact_type, filename, page)
        caption = title
        status = "vision_api_not_configured" if not vision_is_configured(settings) else "vision_analysis_failed"
        if artifact_type == "map" and local_map_reading:
            text = (
                f"{title}. {local_map_reading.get('summary', '')} "
                "Локальный MapReader выделил визуальные признаки карты; мелкие подписи и значения изолиний требуют подключённой Vision/OCR-модели."
            ).strip()
        else:
            text = (
                f"{title}. Визуальный артефакт проиндексирован, но семантическое чтение изображения не выполнено. "
                "Чтобы агент читал легенду, контуры, скважины и подписи на карте, укажите VISION_API_KEY "
                "или DEEPSEEK_OCR_API_KEY в backend/.env."
            )
        metadata = {
            "parser": parser,
            "source": source,
            "visual_status": status,
            "local_visual_type": local_visual_type,
            "local_map_reading": local_map_reading,
            "preview_data_url": preview_url,
        }

    return DocumentArtifact(
        document_name=filename,
        page=page,
        artifact_type=artifact_type,  # type: ignore[arg-type]
        title=str(title),
        caption=str(caption),
        text=text,
        metadata=metadata,
    )


def _merge_visual_artifacts(base: list[DocumentArtifact], visual: list[DocumentArtifact]) -> list[DocumentArtifact]:
    if not visual:
        return base
    visual_keys = {(a.page, a.artifact_type) for a in visual}
    merged = [
        a for a in base
        if not (a.metadata.get("parser") == "caption-detector" and (a.page, a.artifact_type) in visual_keys)
    ]
    seen = {(a.page, a.artifact_type, a.title, a.text[:80]) for a in merged}
    for artifact in visual:
        key = (artifact.page, artifact.artifact_type, artifact.title, artifact.text[:80])
        if key not in seen:
            merged.append(artifact)
            seen.add(key)
    return merged


def _fallback_visual_title(artifact_type: str, filename: str, page: int | None) -> str:
    if artifact_type == "map":
        return "Структурная карта по кровле пласта"
    noun = "Визуальный артефакт"
    suffix = f", стр. {page}" if page else ""
    return f"{noun}: {filename}{suffix}"


def _merge_local_map_reading(payload: dict[str, Any], local_reading: dict[str, Any]) -> dict[str, Any]:
    enriched = dict(payload)
    if not enriched.get("summary") and local_reading.get("summary"):
        enriched["summary"] = local_reading["summary"]
    for key in ("interpretation", "observations", "legend", "contours", "wells", "limitations"):
        base = enriched.get(key)
        base_list = base if isinstance(base, list) else ([base] if base else [])
        add_list = local_reading.get(key)
        add_list = add_list if isinstance(add_list, list) else ([add_list] if add_list else [])
        merged: list[Any] = []
        seen: set[str] = set()
        for item in [*base_list, *add_list]:
            token = str(item)
            if token and token not in seen:
                merged.append(item)
                seen.add(token)
        if merged:
            enriched[key] = merged
    if local_reading.get("confidence") and not enriched.get("confidence"):
        enriched["confidence"] = local_reading["confidence"]
    return enriched


def _local_map_reading(image_bytes: bytes) -> dict[str, Any]:
    """Local, no-API visual reading for map-like pages.

    Enhanced version: uses higher resolution sampling, spatial heatmaps,
    edge/line detection via Pillow filters, blob detection for well points,
    and a legend-region heuristic to separate map body from legend block.
    All of this runs with only Pillow (already a project dependency).
    """
    try:
        from PIL import Image, ImageFilter, ImageEnhance
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")

        # Work at 640px wide for better spatial resolution than the old 520px
        img.thumbnail((640, 640))
        width, height = img.size
        pixels = list(img.get_flattened_data() if hasattr(img, "get_flattened_data") else img.getdata())
        total = max(len(pixels), 1)

        # ── 1. Colour census with bounding boxes ─────────────────────────────
        # Extended palette: orange (Perm/Jurassic), brown (deep), purple (C3),
        # cyan (water contacts), grey (structural lines).
        colour_keys = ("green", "yellow", "orange", "brown", "purple",
                       "blue", "cyan", "red", "dark", "grey", "ink")
        counts: dict[str, int] = {k: 0 for k in colour_keys}
        boxes: dict[str, list[int] | None] = {k: None for k in colour_keys}

        def add_box(key: str, x: int, y: int) -> None:
            b = boxes.get(key)
            if b is None:
                boxes[key] = [x, y, x, y]
            else:
                b[0] = min(b[0], x); b[1] = min(b[1], y)
                b[2] = max(b[2], x); b[3] = max(b[3], y)

        # Spatial grid: divide image into 4×4 zones to detect clustering
        grid_w = max(width // 4, 1)
        grid_h = max(height // 4, 1)
        # zone_colour[zone_idx][colour] = count
        zone_colour: dict[int, dict[str, int]] = {i: {k: 0 for k in colour_keys} for i in range(16)}

        for idx, (r, g, b) in enumerate(pixels):
            x = idx % width
            y = idx // width
            zone = (y // grid_h) * 4 + (x // grid_w)
            zone = min(zone, 15)

            near_white = r > 235 and g > 235 and b > 235
            near_black = r < 40 and g < 40 and b < 40
            sat = max(r, g, b) - min(r, g, b)  # crude saturation

            if not near_white and sat > 20:
                counts["ink"] += 1

            # Green (C1 zones, vegetation-free geological fills)
            if g > 110 and g > r * 1.08 and g > b * 1.1 and not near_white:
                counts["green"] += 1; add_box("green", x, y)
                zone_colour[zone]["green"] += 1

            # Yellow / khaki (C2, sediment layers)
            if r > 160 and g > 140 and b < 140 and sat > 25 and not near_white:
                counts["yellow"] += 1; add_box("yellow", x, y)
                zone_colour[zone]["yellow"] += 1

            # Orange (warm geological fills)
            if r > 180 and g > 100 and g < 175 and b < 100 and not near_white:
                counts["orange"] += 1; add_box("orange", x, y)
                zone_colour[zone]["orange"] += 1

            # Brown / dark-orange
            if r > 120 and g > 60 and g < 110 and b < 80 and not near_white:
                counts["brown"] += 1; add_box("brown", x, y)
                zone_colour[zone]["brown"] += 1

            # Purple / violet (C3, deep formations)
            if r > 100 and b > 100 and r > g * 1.1 and b > g * 1.1 and not near_white:
                counts["purple"] += 1; add_box("purple", x, y)
                zone_colour[zone]["purple"] += 1

            # Blue (water, contacts, boundaries)
            if b > 110 and b > r * 1.2 and b > g * 1.1 and not near_white:
                counts["blue"] += 1; add_box("blue", x, y)
                zone_colour[zone]["blue"] += 1

            # Cyan (GWC / gas contacts)
            if b > 130 and g > 130 and b > r * 1.3 and g > r * 1.3 and not near_white:
                counts["cyan"] += 1; add_box("cyan", x, y)
                zone_colour[zone]["cyan"] += 1

            # Red (wells, faults, important markers)
            if r > 140 and r > g * 1.4 and r > b * 1.35 and not near_white:
                counts["red"] += 1; add_box("red", x, y)
                zone_colour[zone]["red"] += 1

            # Dark (text, isolines, structural lines)
            if r < 90 and g < 90 and b < 90 and not near_black:
                counts["dark"] += 1; add_box("dark", x, y)
                zone_colour[zone]["dark"] += 1

            # Grey (administrative boundaries, grid lines)
            if 80 < r < 185 and abs(r - g) < 18 and abs(g - b) < 18 and sat < 22 and not near_white:
                counts["grey"] += 1; add_box("grey", x, y)
                zone_colour[zone]["grey"] += 1

        def ratio(key: str) -> float:
            return counts[key] / total

        def pos_from_box(box: list[int] | None) -> str:
            if not box:
                return ""
            cx = (box[0] + box[2]) / 2 / max(width, 1)
            cy = (box[1] + box[3]) / 2 / max(height, 1)
            horiz = "западная часть" if cx < .33 else "восточная часть" if cx > .67 else "центральная часть"
            vert = "северная" if cy < .33 else "южная" if cy > .67 else "центральная"
            return f"{vert}, {horiz}"

        # ── 2. Edge / line density (isoline detection) ────────────────────────
        # Apply Canny-like detection using Pillow's FIND_EDGES filter on a
        # greyscale version of the map body (exclude legend region heuristic).
        map_body_right = int(width * 0.82)  # legend usually in right 18%
        try:
            grey = img.crop((0, 0, map_body_right, height)).convert("L")
            edges = grey.filter(ImageFilter.FIND_EDGES)
            edge_pixels = list(edges.get_flattened_data() if hasattr(edges, "get_flattened_data") else edges.getdata())
            edge_total = max(len(edge_pixels), 1)
            edge_density = sum(1 for p in edge_pixels if p > 30) / edge_total
            has_isolines = edge_density > 0.04  # >4% edge pixels = complex line network
            has_dense_lines = edge_density > 0.09
        except Exception:
            edge_density = 0.0
            has_isolines = ratio("dark") > 0.012
            has_dense_lines = False

        # ── 3. Legend block detection (right strip) ───────────────────────────
        # If the rightmost 20% of the image has a very different colour profile
        # than the map body, it's likely the legend block.
        legend_strip_x = int(width * 0.80)
        legend_strip_pixels = [
            pixels[y * width + x]
            for y in range(height)
            for x in range(legend_strip_x, width)
            if y * width + x < len(pixels)
        ]
        legend_white_ratio = (
            sum(1 for r, g, b in legend_strip_pixels if r > 235 and g > 235 and b > 235)
            / max(len(legend_strip_pixels), 1)
        )
        has_legend_block = legend_white_ratio > 0.45  # legend area is mostly white paper

        # ── 4. Well point detection ───────────────────────────────────────────
        # Count isolated red/dark clusters in the map body zone.
        # A simple proxy: count 8×8 tiles with significant red/dark pixels.
        tile_size = 8
        well_tile_count = 0
        for ty in range(0, height, tile_size):
            for tx in range(0, map_body_right, tile_size):
                tile_red = 0
                for dy in range(tile_size):
                    for dx in range(tile_size):
                        px = tx + dx; py = ty + dy
                        if px < width and py < height:
                            idx2 = py * width + px
                            if idx2 < len(pixels):
                                r, g, b = pixels[idx2]
                                if r > 130 and r > g * 1.35 and r > b * 1.3:
                                    tile_red += 1
                if tile_red >= 3:
                    well_tile_count += 1
        # Normalise by total tiles
        total_tiles = max((height // tile_size) * (map_body_right // tile_size), 1)
        well_tile_ratio = well_tile_count / total_tiles

        # ── 5. Assemble observations ──────────────────────────────────────────
        observations: list[str] = []
        legend: list[dict[str, str]] = []
        contours: list[dict[str, str]] = []
        wells: list[dict[str, str]] = []
        interpretation: list[str] = []

        # Colour zone observations (more nuanced than before)
        colour_specs = [
            ("yellow",  0.005, "жёлтая",   "категория C1 или основная площадная зона", "C1 или основная зона"),
            ("green",   0.003, "зелёная",   "категория C2 или внутренняя зона",          "C2 или внутренняя зона"),
            ("orange",  0.003, "оранжевая", "пермские или юрские отложения",             "пермские/юрские отложения"),
            ("brown",   0.002, "коричневая","глубокие горизонты или доюрский фундамент", "доюрский фундамент"),
            ("purple",  0.002, "фиолетовая","категория C3 или переходная зона",          "категория C3"),
        ]
        for ckey, thresh, cname, meaning, interp in colour_specs:
            if ratio(ckey) > thresh:
                pos = pos_from_box(boxes[ckey])
                observations.append(
                    f"{cname} заливка ({pos}) — {meaning} (покрывает {ratio(ckey)*100:.1f}% площади)"
                )
                legend.append({"label": f"{cname} зона", "symbol": "площадная заливка", "meaning": meaning})
                interpretation.append(f"{cname} заливка — {interp}")

        # Contour / boundary observations
        if ratio("blue") > 0.001:
            pos = pos_from_box(boxes["blue"])
            label = "ВНК или внешний контур нефтегазоносности" if ratio("blue") > 0.005 else "граница лицензионного участка или административная граница"
            contours.append({"label": label, "style": "синяя линия", "confidence": "medium"})
            observations.append(f"синие линии ({pos}) — {label}")
            interpretation.append(f"синие линии — {label}")

        if ratio("cyan") > 0.001:
            contours.append({"label": "ГВК или газоводяной контакт", "style": "голубая линия", "confidence": "medium"})
            observations.append("голубые/циановые линии — возможный ГВК или водяной контакт")

        if has_isolines:
            density_desc = "плотная сеть" if has_dense_lines else "умеренная сеть"
            contours.append({
                "label": "изогипсы/изолинии",
                "style": "тонкие тёмные линии",
                "confidence": "high" if has_dense_lines else "medium",
            })
            observations.append(
                f"обнаружена {density_desc} изолиний (плотность рёбер {edge_density*100:.1f}%) — характерно для структурной карты"
            )
            interpretation.append("тонкие изолинии — изогипсы структурной карты")

        if ratio("grey") > 0.008:
            contours.append({"label": "административные/лицензионные границы", "style": "серая линия", "confidence": "low"})
            observations.append("серые линии — возможные административные границы или сетка координат")

        # Well point observations
        if well_tile_ratio > 0.02 or ratio("red") > 0.0003:
            n_approx = max(int(well_tile_count * 0.6), 1)  # rough estimate
            confidence = "medium" if well_tile_ratio > 0.04 else "low"
            wells.append({
                "id": f"~{n_approx} точечных маркеров",
                "type": "unknown",
                "relative_position": pos_from_box(boxes["red"]) or "распределены по площади",
                "confidence": confidence,
            })
            observations.append(
                f"обнаружено ≈{n_approx} точечных маркеров (красные/тёмные кластеры) — вероятные скважины"
            )
            interpretation.append(f"≈{n_approx} точечных маркеров — вероятные скважины")

        if has_legend_block:
            observations.append("в правой части изображения выделяется блок условных обозначений (легенда карты)")

        # ── 6. Summary ───────────────────────────────────────────────────────
        zone_count = sum(1 for k in ("yellow", "green", "orange", "brown", "purple") if ratio(k) > 0.002)
        has_contacts = ratio("blue") > 0.001 or ratio("cyan") > 0.001
        well_desc = f", ≈{int(well_tile_count*0.6)} точечных маркеров скважин" if well_tile_ratio > 0.02 else ""
        contact_desc = ", линии контактов/контуров" if has_contacts else ""
        summary = (
            f"Карта содержит {zone_count} цветовых зон{contact_desc}{well_desc}. "
            f"Плотность линий: {edge_density*100:.1f}% — {'характерна для изогипс' if has_isolines else 'ниже порога изогипс'}. "
            "Точные числа изолиний и номера скважин требуют Vision/OCR."
        )

        limitations = [
            "точные значения изолиний (числа вдоль линий) не считаны — нужен Vision/OCR",
            "номера скважин не читаются без OCR",
        ]
        if not has_isolines:
            limitations.append("изолинии не обнаружены локальным детектором — возможен слабый контраст скана")

        confidence = "medium" if (observations and (has_isolines or well_tile_ratio > 0.02)) else "low"
        return {
            "summary": summary,
            "interpretation": interpretation,
            "observations": observations,
            "legend": legend,
            "contours": contours,
            "wells": wells,
            "limitations": limitations,
            "confidence": confidence,
            "_local_stats": {
                "edge_density": round(edge_density, 4),
                "well_tiles": well_tile_count,
                "has_legend_block": has_legend_block,
            },
        }
    except Exception:
        return {}


def _make_preview_data_url(image_bytes: bytes, mime_type: str) -> str:
    # Keep prototype self-contained: previews are returned inline in metadata.
    # For production, store previews in object storage and return a URL instead.
    # PDF renders often contain a tiny map in the middle of a large white page.
    # Crop near-white margins for the preview only, so the UI shows the actual map.
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        img = _crop_near_white_margins(img)
        img.thumbnail((1400, 1000))
        out = io.BytesIO()
        img.save(out, format="JPEG", quality=84, optimize=True)
        return image_data_url(out.getvalue(), "image/jpeg")
    except Exception:
        logger.debug("Preview crop/resize failed; returning inline preview")
        max_bytes = 900_000
        if len(image_bytes) <= max_bytes:
            return image_data_url(image_bytes, mime_type)
        return image_data_url(image_bytes[:max_bytes], mime_type)


def _crop_near_white_margins(img):
    """Crop empty page margins while keeping a small visual border.

    The algorithm is intentionally simple and dependency-light: it finds pixels
    that are not near-white/near-background and crops to their bounding box.
    It affects only thumbnails returned to the frontend, not the original image
    sent to the vision model.
    """
    width, height = img.size
    if width < 40 or height < 40:
        return img
    small = img.copy()
    scale = max(width / 900, height / 900, 1)
    if scale > 1:
        small.thumbnail((int(width / scale), int(height / scale)))
    sw, sh = small.size
    pixels = small.load()
    xs: list[int] = []
    ys: list[int] = []
    for y in range(sh):
        for x in range(sw):
            r, g, b = pixels[x, y]
            # Keep dark text/lines and colored geology zones. Ignore white paper.
            colored = max(r, g, b) - min(r, g, b) > 18 and max(r, g, b) < 252
            dark = min(r, g, b) < 220
            if colored or dark:
                xs.append(x)
                ys.append(y)
    if not xs or not ys:
        return img
    left, right = min(xs), max(xs)
    top, bottom = min(ys), max(ys)
    # Avoid accidental crop when the whole page is actually filled.
    if (right - left) / sw > 0.92 and (bottom - top) / sh > 0.92:
        return img
    margin = max(12, int(min(sw, sh) * 0.025))
    left = max(0, left - margin)
    top = max(0, top - margin)
    right = min(sw - 1, right + margin)
    bottom = min(sh - 1, bottom + margin)
    inv = width / sw
    return img.crop((
        max(0, int(left * inv)),
        max(0, int(top * inv)),
        min(width, int((right + 1) * inv)),
        min(height, int((bottom + 1) * inv)),
    ))


def _guess_visual_type(image_bytes: bytes) -> str:
    """Cheap local heuristic: geology maps often have colored zones + contour lines.

    Enhanced: also checks edge density (isolines) and aspect ratio.
    Map signals: coloured zones + line network OR dominant single colour fill.
    """
    try:
        from PIL import Image, ImageFilter
        img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        img.thumbnail((380, 380))
        pixels = list(img.get_flattened_data() if hasattr(img, "get_flattened_data") else img.getdata())
        total = max(len(pixels), 1)

        colored = blue = green = yellow = red = orange = purple = 0
        for r, g, b in pixels:
            near_white = r > 235 and g > 235 and b > 235
            sat = max(r, g, b) - min(r, g, b)
            if not near_white and sat > 30 and max(r, g, b) < 250:
                colored += 1
            if b > 110 and b > r * 1.2 and b > g * 1.1 and not near_white:
                blue += 1
            if g > 110 and g > r * 1.08 and g > b * 1.1 and not near_white:
                green += 1
            if r > 155 and g > 135 and b < 140 and not near_white:
                yellow += 1
            if r > 130 and r > g * 1.35 and r > b * 1.3 and not near_white:
                red += 1
            if r > 175 and g > 95 and g < 175 and b < 100 and not near_white:
                orange += 1
            if r > 100 and b > 100 and r > g * 1.1 and b > g * 1.1 and not near_white:
                purple += 1

        color_ratio = colored / total
        fill_signal = (blue + green + yellow + orange + purple) / total
        point_signal = red / total

        # Edge density for isoline detection
        grey = img.convert("L")
        edges = grey.filter(ImageFilter.FIND_EDGES)
        edge_pixels = list(edges.get_flattened_data() if hasattr(edges, "get_flattened_data") else edges.getdata())
        edge_density = sum(1 for p in edge_pixels if p > 25) / max(len(edge_pixels), 1)

        # Map: strong colour fill OR (moderate colour + dense lines) OR (isolated red points + colour)
        if fill_signal > 0.02 and edge_density > 0.035:
            return "map"
        if color_ratio > 0.03 and fill_signal > 0.012:
            return "map"
        if point_signal > 0.0005 and fill_signal > 0.005:
            return "map"
    except Exception:
        pass
    return "figure"

def _strip_code_fence(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    return text.strip()


def _safe_int(value: Any) -> int | None:
    try:
        return int(value)
    except Exception:
        return None
